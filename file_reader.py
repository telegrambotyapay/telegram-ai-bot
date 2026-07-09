"""
Yüklenen dosyaları (PDF, Word, Excel, CSV, metin) okuyup düz metne çevirme.
"""
import io
import csv
import logging

from pypdf import PdfReader
from docx import Document
import openpyxl

logger = logging.getLogger(__name__)


class FileReadError(Exception):
    """Dosya okunamadığında fırlatılır."""


def read_pdf(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        texts = [page.extract_text() or "" for page in reader.pages]
        content = "\n".join(texts).strip()
        if not content:
            raise FileReadError(
                "PDF'den metin çıkarılamadı (taranmış/görsel bir PDF olabilir, "
                "bu durumda fotoğraf olarak göndermeyi dene)."
            )
        return content
    except FileReadError:
        raise
    except Exception as e:
        raise FileReadError(f"PDF okunamadı: {e}") from e


def read_docx(file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        table_lines = []
        for table in doc.tables:
            for row in table.rows:
                table_lines.append(" | ".join(cell.text for cell in row.cells))
        content = "\n".join(paragraphs + table_lines).strip()
        if not content:
            raise FileReadError("Word belgesinden metin çıkarılamadı.")
        return content
    except FileReadError:
        raise
    except Exception as e:
        raise FileReadError(f"Word belgesi okunamadı: {e}") from e


def read_xlsx(file_bytes: bytes) -> str:
    try:
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        lines = []
        for sheet in wb.worksheets:
            lines.append(f"--- {sheet.title} ---")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    lines.append(" | ".join(cells))
        content = "\n".join(lines).strip()
        if not content:
            raise FileReadError("Excel dosyasından veri okunamadı.")
        return content
    except FileReadError:
        raise
    except Exception as e:
        raise FileReadError(f"Excel dosyası okunamadı: {e}") from e


def read_csv(file_bytes: bytes) -> str:
    try:
        text = file_bytes.decode("utf-8", errors="replace")
        rows = list(csv.reader(io.StringIO(text)))
        content = "\n".join(" | ".join(row) for row in rows).strip()
        if not content:
            raise FileReadError("CSV dosyasından veri okunamadı.")
        return content
    except FileReadError:
        raise
    except Exception as e:
        raise FileReadError(f"CSV dosyası okunamadı: {e}") from e


def read_txt(file_bytes: bytes) -> str:
    try:
        content = file_bytes.decode("utf-8", errors="replace").strip()
        if not content:
            raise FileReadError("Metin dosyası boş.")
        return content
    except FileReadError:
        raise
    except Exception as e:
        raise FileReadError(f"Metin dosyası okunamadı: {e}") from e


def read_file(filename: str, file_bytes: bytes) -> str:
    """Dosya uzantısına göre uygun okuyucuyu seçer."""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext == "pdf":
        return read_pdf(file_bytes)
    if ext == "docx":
        return read_docx(file_bytes)
    if ext in ("xlsx", "xlsm"):
        return read_xlsx(file_bytes)
    if ext == "csv":
        return read_csv(file_bytes)
    if ext in ("txt", "md"):
        return read_txt(file_bytes)
    raise FileReadError(
        f"'.{ext}' uzantılı dosyaları henüz desteklemiyorum. "
        "Desteklenenler: PDF, DOCX, XLSX, CSV, TXT."
    )
"""
Yüklenen dosyaları (PDF, Word, Excel, CSV, metin) okuyup düz metne çevirme.
"""
import io
import csv
import logging

from pypdf import PdfReader
from docx import Document
import openpyxl

logger = logging.getLogger(__name__)


class FileReadError(Exception):
    """Dosya okunamadığında fırlatılır."""


def read_pdf(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        texts = [page.extract_text() or "" for page in reader.pages]
        content = "\n".join(texts).strip()
        if not content:
            raise FileReadError(
                "PDF'den metin çıkarılamadı (taranmış/görsel bir PDF olabilir, "
                "bu durumda fotoğraf olarak göndermeyi dene)."
            )
        return content
    except FileReadError:
        raise
    except Exception as e:
        raise FileReadError(f"PDF okunamadı: {e}") from e


def read_docx(file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        table_lines = []
        for table in doc.tables:
            for row in table.rows:
                table_lines.append(" | ".join(cell.text for cell in row.cells))
        content = "\n".join(paragraphs + table_lines).strip()
        if not content:
            raise FileReadError("Word belgesinden metin çıkarılamadı.")
        return content
    except FileReadError:
        raise
    except Exception as e:
        raise FileReadError(f"Word belgesi okunamadı: {e}") from e


def read_xlsx(file_bytes: bytes) -> str:
    try:
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        lines = []
        for sheet in wb.worksheets:
            lines.append(f"--- {sheet.title} ---")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    lines.append(" | ".join(cells))
        content = "\n".join(lines).strip()
        if not content:
            raise FileReadError("Excel dosyasından veri okunamadı.")
        return content
    except FileReadError:
        raise
    except Exception as e:
        raise FileReadError(f"Excel dosyası okunamadı: {e}") from e


def read_csv(file_bytes: bytes) -> str:
    try:
        text = file_bytes.decode("utf-8", errors="replace")
        rows = list(csv.reader(io.StringIO(text)))
        content = "\n".join(" | ".join(row) for row in rows).strip()
        if not content:
            raise FileReadError("CSV dosyasından veri okunamadı.")
        return content
    except FileReadError:
        raise
    except Exception as e:
        raise FileReadError(f"CSV dosyası okunamadı: {e}") from e


def read_txt(file_bytes: bytes) -> str:
    try:
        content = file_bytes.decode("utf-8", errors="replace").strip()
        if not content:
            raise FileReadError("Metin dosyası boş.")
        return content
    except FileReadError:
        raise
    except Exception as e:
        raise FileReadError(f"Metin dosyası okunamadı: {e}") from e


def read_file(filename: str, file_bytes: bytes) -> str:
    """Dosya uzantısına göre uygun okuyucuyu seçer."""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext == "pdf":
        return read_pdf(file_bytes)
    if ext == "docx":
        return read_docx(file_bytes)
    if ext in ("xlsx", "xlsm"):
        return read_xlsx(file_bytes)
    if ext == "csv":
        return read_csv(file_bytes)
    if ext in ("txt", "md"):
        return read_txt(file_bytes)
    raise FileReadError(
        f"'.{ext}' uzantılı dosyaları henüz desteklemiyorum. "
        "Desteklenenler: PDF, DOCX, XLSX, CSV, TXT."
    )
