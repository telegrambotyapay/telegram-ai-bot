"""
Analiz metnini Word (.docx), Excel (.xlsx) ya da PDF dosyasına aktarma.
"""
import re
import tempfile

from docx import Document
import openpyxl
from fpdf import FPDF


def create_docx(title: str, content: str) -> str:
    doc = Document()
    doc.add_heading(title, level=1)
    for line in content.split("\n"):
        line = line.strip()
        if line:
            doc.add_paragraph(line)
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp.close()
    doc.save(tmp.name)
    return tmp.name


def _split_row(line: str):
    if "|" in line:
        parts = [c.strip() for c in line.split("|") if c.strip()]
        if parts:
            return parts
    parts = re.split(r"\s{2,}|\t", line.strip())
    parts = [p for p in parts if p]
    return parts if len(parts) > 1 else [line.strip()]


def create_xlsx(title: str, content: str) -> str:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = (title or "Analiz")[:31]
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue
        ws.append(_split_row(line))
    tmp = tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False)
    tmp.close()
    wb.save(tmp.name)
    return tmp.name


def _break_long_words(text: str, max_word_len: int = 60) -> str:
    """
    Boşluk içermeyen çok uzun diziler (uzun linkler, hash'ler vb.) PDF satır
    kaydırma motorunu çökertebiliyor ("Not enough horizontal space..." hatası).
    Bu tür kelimeleri belirli aralıklarla bölerek bu hatayı önlüyoruz.
    """
    words = text.split(" ")
    fixed_words = []
    for word in words:
        if len(word) > max_word_len:
            chunks = [word[i:i + max_word_len] for i in range(0, len(word), max_word_len)]
            fixed_words.append(" ".join(chunks))
        else:
            fixed_words.append(word)
    return " ".join(fixed_words)


def create_pdf(title: str, content: str) -> str:
    """Metni başlıklı bir PDF'e yazar, geçici dosya yolunu döner."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Türkçe karakterleri destekleyen bir yazı tipi bulunamazsa, güvenli
    # (ASCII'ye yakın) bir dönüştürme yaparak en azından çökmesini önlüyoruz.
    def safe(text: str) -> str:
        return text.encode("latin-1", errors="replace").decode("latin-1")

    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(0, 10, safe(_break_long_words(title)))
    pdf.ln(4)

    pdf.set_font("Helvetica", "", 11)
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            pdf.ln(3)
            continue
        try:
            pdf.multi_cell(0, 7, safe(_break_long_words(line)))
        except Exception:
            # Bu satır her ne sebeple olursa olsun render edilemedi,
            # tüm PDF'i çökertmek yerine kısaltılmış haliyle devam et.
            try:
                pdf.multi_cell(0, 7, safe(_break_long_words(line))[:80] + " [...]")
            except Exception:
                pass

    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    tmp.close()
    pdf.output(tmp.name)
    return tmp.name
